import re
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.io import write_image
from docx.shared import Inches


def extract_values_month(text):
    # Extracts month-value-unit triplets from text.
    # Supports both bullet/colon formats and Markdown-style table rows.
    pattern1 = r"[-•]?\s*([A-Za-z]+)\s*[:=\-]\s*([-+]?\d*\.?\d+)\s*(mm|mbar|°C)?"
    data1 = [(m, float(v), u or "") for m, v, u in re.findall(pattern1, text)]

    pattern2 = r"\|\s*([A-Za-z]+)\s*\|\s*([-+]?\d*\.?\d+)\s*(?:\|.*)?"
    data2 = [(m, float(v), "") for m, v in re.findall(pattern2, text)]

    # Combine both results
    return data1 + data2


def extract_values_annual(text):
    # Extracts (year, value, unit) tuples from the text (annual format).
    pattern = r"(?:Year\s*)?(\d{4})\s*[:=\-]\s*([-+]?\d*\.?\d+)\s*(mm|mbar|°C)?"
    return [(y, float(v), u) for y, v, u in re.findall(pattern, text)]


# -------------------------------
# Classification of 10 known question types
# -------------------------------
def question_classification(question: str) -> str:
    # Identifies the question type to determine which chart to generate.
    question = question.lower()
    if "highest value" in question:
        return "barchart"
    if "lowest value" in question:
        return "barchart"
    if "discrepancy" in question:
        return "barchart_discrepancy"
    if "average annual" in question:
        return "media_annuale"
    if "median" in question:
        return "boxplot"
    if "mode" in question:
        return "monovalore"
    if "differ" in question and "maximum and minimum" in question:
        return "barchart_2columns"
    if "not recorded" in question:
        return "monovalore"
    if "drop below" in question:
        return "monovalore"
    if "changed over the last two years" in question:
        return "barchart_averageannual_diff"
    return "altro"


def graph_from_df(df, question):
    # Creates a bar chart automatically from a DataFrame, adapting based on the question.
    df2 = df.copy()

    # Identify X column (month or fallback to first column)
    x_col = None
    for col in df2.columns:
        if "month" in col.lower():
            x_col = col
            break
    if x_col is None:
        x_col = df2.columns[0]  # fallback to first column

    df2[x_col] = df2[x_col].astype(str).str.strip()

    # Detect numeric columns
    numeric_cols = df2.select_dtypes(include=["number"]).columns.tolist()
    if not numeric_cols:
        # Try to coerce text columns to numbers using regex
        for col in df2.columns:
            try:
                df2[col] = pd.to_numeric(
                    df2[col].astype(str).str.replace(r"[^\d\.\-]", "", regex=True),
                    errors='coerce'
                )
            except Exception:
                continue
        numeric_cols = df2.select_dtypes(include=["number"]).columns.tolist()

    if not numeric_cols:
        return None  # no numeric column found

    # Select y-axis column based on keywords in the question
    question_lower = question.lower()
    if "min" in question_lower:
        y_col = next((c for c in numeric_cols if "min" in c.lower()), numeric_cols[0])
    elif "max" in question_lower:
        y_col = next((c for c in numeric_cols if "max" in c.lower()), numeric_cols[0])
    elif "mean" in question_lower or "average" in question_lower:
        y_col = next((c for c in numeric_cols if "mean" in c.lower()), numeric_cols[0])
    else:
        possible_cols = [
            c for c in numeric_cols if any(k in str(c).lower() for k in ["max", "min", "mean", "temperature", "precipitation", "pressure"])
        ]
        y_col = possible_cols[0] if possible_cols else numeric_cols[0]

    # Drop rows with missing values
    df2 = df2[[x_col, y_col]].dropna()
    if df2.empty:
        return None

    # Create the bar plot
    fig = px.bar(
        df2,
        x=x_col,
        y=y_col,
        labels={x_col: "Month" if "month" in str(x_col).lower() else x_col, y_col: "Value"},
        title=f"{y_col} - {question}"
    )

    # Highlight min/max bars in red
    yname = str(y_col).lower()
    if "min" in yname:
        pos_idx = df2[y_col].idxmin()
        colors = ["skyblue"] * len(df2)
        pos = list(df2.index).index(pos_idx)
        colors[pos] = "red"
        fig.data[0].marker.color = colors
    elif "max" in yname:
        pos_idx = df2[y_col].idxmax()
        colors = ["skyblue"] * len(df2)
        pos = list(df2.index).index(pos_idx)
        colors[pos] = "red"
        fig.data[0].marker.color = colors
    else:
        fig.update_traces(marker_color="skyblue")

    # Rotate labels if too many categories
    if len(df2) > 8:
        fig.update_layout(xaxis_tickangle=-45)

    # Auto-adjust Y-axis with margin
    ymin = df2[y_col].min()
    ymax = df2[y_col].max()
    margin = 0.05 * (ymax - ymin) if ymin != ymax else 0.05 * abs(ymin) if ymin != 0 else 1
    fig.update_layout(yaxis=dict(range=[ymin - margin, ymax + margin]))

    # Add text labels above bars
    fig.update_traces(text=df2[y_col], textposition="outside")

    return fig


def save_graph(fig, question):
    # Saves Plotly figure as a PNG image and returns the file path.
    chart_path = f"grafico_{abs(hash(question))}.png"
    fig.update_layout(margin=dict(t=100, b=50, l=80, r=80))  # avoid cutting title
    fig.write_image(chart_path, width=1200, height=500, scale=2)
    return chart_path


def generate_graphics(doc, chat_history, dfs=None):
    """
    Generates graphs from chat_history and inserts them into a Word document.
    """

    # Normalize DataFrames input
    if dfs is None:
        dfs = []
    dfs = [pd.DataFrame(df) for df in dfs if df is not None and not df.empty]

    for entry in chat_history:
        question = entry["question"]
        answer = entry["answer"]
        type = question_classification(question)
        
        chart_created = False
        chart_path = None

        # ------------------- Boxplot (for median questions) -------------------
        if type == "boxplot":
            doc.add_heading("Data Charts", level=1)
            data = extract_values_month(answer)
       
            if not data:
                data = extract_values_annual(answer)

            if data:
                labels, values, unit = zip(*data)
                values = [float(v) for v in values]
                unit_y = unit[0] if unit else ""

                fig = go.Figure()
                fig.add_trace(go.Box(
                    y=values,
                    boxpoints="all",  # show all points
                    jitter=0.5,
                    pointpos=-1.8,
                    name="Data",
                    marker_color='skyblue'
                ))

                fig.update_layout(
                    title=question,
                    yaxis_title=f"Value ({unit_y})"
                )

                chart_path = save_graph(fig, question)
                chart_created = True
        

        # ------------------- Max vs Min Comparison -------------------
        elif type == "barchart_2columns":
            doc.add_heading("Data Charts", level=1)
            year = extract_values_annual(answer)

            val_max = val_min = diff_val = None
            unit = ""
            source = None

            # 1) Try extracting from DataFrame
            if "df" in locals() and not df.empty:
                cols = [c.lower().replace(" ", "_") for c in df.columns]
                col_map = dict(zip(cols, df.columns))

                max_col = next((col_map[c] for c in cols if "max" in c), None)
                min_col = next((col_map[c] for c in cols if "min" in c), None)

                if max_col and min_col:
                    try:
                        val_max = df[max_col].max()
                        val_min = df[min_col].min()
                        diff_val = val_max - val_min
                        source = "dataframe"
                    except Exception as e:
                        print(f"[barchart_2columns] DataFrame error: {e}")
                        val_max = val_min = diff_val = None

            # 2) Regex-based fallback
            if val_max is None or val_min is None:
                pattern = r"Maximum .*?[:]\s*([-+]?\d*\.\d+)\s*(°C|mm|mbar)?.*?Minimum .*?[:]\s*([-+]?\d*\.\d+)\s*(°C|mm|mbar)?"
                m = re.search(pattern, answer, re.IGNORECASE | re.DOTALL)
                if m:
                    val_max = float(m.group(1))
                    val_min = float(m.group(3))
                    unit = m.group(2) or m.group(4) or ""
                    diff_val = val_max - val_min
                    source = "regex"
                else:
                    # Separate regex fallback
                    max_pat = r"maximum(?:\s+temperature|\s+pressure|\s+precipitation)?[^:=\n\r]*[:=\-–]\s*([-+]?\d*\.\d+)\s*(°C|mm|mbar)?"
                    min_pat = r"minimum(?:\s+temperature|\s+pressure|\s+precipitation)?[^:=\n\r]*[:=\-–]\s*([-+]?\d*\.\d+)\s*(°C|mm|mbar)?"
                    max_m = re.search(max_pat, answer, re.IGNORECASE | re.DOTALL)
                    min_m = re.search(min_pat, answer, re.IGNORECASE | re.DOTALL)
                    if max_m and min_m:
                        val_max = float(max_m.group(1))
                        val_min = float(min_m.group(1))
                        unit = (max_m.group(2) or min_m.group(2) or "").strip()
                        diff_val = val_max - val_min
                        source = "regex"

            # 3) Fallback: use all numeric values if necessary
            if val_max is None or val_min is None:
                numbers = [float(x) for x in re.findall(r"\d+\.\d+", answer)]
                if numbers:
                    val_max, val_min = max(numbers), min(numbers)
                    diff_val = val_max - val_min
                    unit = ""
                    source = "numbers"

            # 4) Create bar chart if valid
            if val_max is not None and val_min is not None:
                plotted_min = min(val_min, val_max)
                plotted_max = max(val_min, val_max)
                diff_val = plotted_max - plotted_min

                fig = go.Figure(data=[
                    go.Bar(
                        x=["Min", "Max"],
                        y=[plotted_min, plotted_max],
                        marker_color=["steelblue", "crimson"],
                        text=[f"{plotted_min}", f"{plotted_max}"],
                        textposition="outside"
                    )
                ])

                fig.add_annotation(
                    x=0.5, xref="paper",
                    y=plotted_max,
                    text=f"Δ = {plotted_max - plotted_min:.2f} {unit}".strip(),
                    showarrow=False,
                    bgcolor="lightyellow",
                    font=dict(size=14)
                )

                fig.update_layout(
                    title=question,
                    yaxis_title=f"Value ({unit})" if unit else "Value",
                    xaxis_title="",
                    bargap=0.5
                )

                chart_path = save_graph(fig, question)
                chart_created = True
                print(f"[barchart_2columns] Chart created from {source}")
            else:
                print("[barchart_2columns] No valid values found.")


# ------------------- Annual average difference -------------------
        elif type == "barchart_averageannual_diff":
            mean_vals = []
            years = []
            unit = ""

            # Extract last lines (usually contain numeric data)
            last_lines = "\n".join(answer.splitlines()[-10:])

            # Regex to match phrases like:
            # "decreased from 1015.72 mbar in 2017 to 1014.40 mbar in 2018"
            pattern = re.compile(
                r"(?P<val1>[-+]?\d*\.\d+)\s*(?P<unit1>mbar|°C|mm)?(?:\s*(?:in|for)?\s*(?P<year1>20\d{2}))?\s*(?:-|to)\s*(?P<val2>[-+]?\d*\.\d+)\s*(?P<unit2>mbar|°C|mm)?(?:\s*(?:in|for)?\s*(?P<year2>20\d{2}))?",
                re.IGNORECASE
            )

            match = pattern.search(last_lines)
            if match:
                val1 = float(match.group("val1"))
                val2 = float(match.group("val2"))
                unit = match.group("unit1") or match.group("unit2") or ""
                mean_vals = [val1, val2]
                year1 = match.group("year1") or "Year 1"
                year2 = match.group("year2") or "Year 2"
                years = [year1, year2]

            # Fallback using DataFrames
            if len(mean_vals) < 2 and dfs:
                mean_vals = []
                years = []
                for i, df in enumerate(dfs[:2]):
                    numeric_cols = df.select_dtypes(include="number").columns.tolist()
                    if numeric_cols:
                        mean_val = df[numeric_cols[0]].mean()
                        mean_vals.append(mean_val)
                        years.append(f"Year {i+1}")
                        col_name = numeric_cols[0].lower()
                        if "°c" in col_name:
                            unit = "°C"
                        elif "mm" in col_name:
                            unit = "mm"
                        elif "mbar" in col_name:
                            unit = "mbar"

            # Create chart if 2 values available
            if len(mean_vals) >= 2:
                doc.add_heading("Data Charts", level=1)
                val1, val2 = mean_vals[:2]
                year1, year2 = years[:2]
                delta = val2 - val1

                delta_text = f"Δ = {'+' if delta >= 0 else '-'}{abs(delta):.2f} {unit}"

                fig = go.Figure(data=[
                    go.Bar(
                        x=[year1, year2],
                        y=[val1, val2],
                        text=[f"{val1:.3f}", f"{val2:.3f}"],
                        textposition="outside",
                        marker_color=["steelblue", "crimson"]
                    )
                ])

                fig.add_annotation(
                    x=0.5, xref="paper",
                    y=max(val1, val2),
                    text=delta_text,
                    showarrow=False,
                    bgcolor="lightyellow",
                    font=dict(size=14)
                )

                fig.update_layout(
                    title=question,
                    yaxis_title=f"Average ({unit})" if unit else "Average",
                    xaxis_title="Year",
                    bargap=0.5
                )

                chart_path = save_graph(fig, question)
                inline_shape = doc.add_picture(chart_path, width=Inches(6))


# ------------------- Monthly discrepancy -------------------
        elif type == "barchart_discrepancy":
            for df in dfs:
                month_cols = [c for c in df.columns if "month" in c.lower()]
                if not month_cols:
                    continue
                month_col = month_cols[0]

                diff_cols = [c for c in df.columns if any(k in c.lower() for k in ["difference", "discrepancy"])]
                values_df = None
                value_col = None

                if diff_cols:
                    value_col = diff_cols[0]
                    values_df = df[[month_col, value_col]].dropna()

                    if values_df.empty or values_df[value_col].dropna().empty:
                        print(f"[barchart_discrepancy] Column {value_col} empty, trying manual calc")
                        values_df = None

                # Fallback: compute as Max - Min
                if values_df is None:
                    max_cols = [c for c in df.columns if "max" in c.lower()]
                    min_cols = [c for c in df.columns if "min" in c.lower()]
                    if max_cols and min_cols:
                        value_col = "Difference"
                        values_df = df[[month_col, max_cols[0], min_cols[0]]].copy()
                        values_df[value_col] = values_df[max_cols[0]] - values_df[min_cols[0]]
                        values_df = values_df[[month_col, value_col]].dropna()
                    else:
                        print("[barchart_discrepancy] No Difference or Max/Min columns found, skipping.")
                        continue

                if values_df.empty or values_df[value_col].dropna().empty:
                    print("[barchart_discrepancy] No valid data, skipping chart.")
                    continue

                # Highlight maximum
                max_idx = values_df[value_col].idxmax()
                max_pos = values_df.index.get_loc(max_idx)
                colors = ["skyblue"] * len(values_df)
                colors[max_pos] = "crimson"

                # Extract unit if found in text
                unit_match = re.search(
                    r"Average\s+[\w\s]+?\s+(?:for|in)\s+\d{4}\s*=\s*[-+]?\d+(?:\.\d+)?\s*(°C|mm|mbar)?",
                    answer, re.IGNORECASE
                )
                unit = unit_match.group(1) if unit_match else ""

                # Build chart
                fig = px.bar(
                    values_df,
                    x=month_col,
                    y=value_col,
                    text=value_col,
                    title=f"{value_col} per {month_col}",
                    labels={value_col: f"{value_col} ({unit})" if unit else value_col}
                )
                fig.update_traces(marker_color=colors, textposition="outside")
                if len(values_df) > 8:
                    fig.update_layout(xaxis_tickangle=-45)

                chart_path = save_graph(fig, question)
                chart_created = True


        # ------------------- Generic bar chart -------------------
        elif type == "barchart":
            doc.add_heading("Data Charts", level=1)
            for df in dfs:
                fig = graph_from_df(df, question)
                if fig:
                    chart_path = save_graph(fig, question)
                    chart_created = True

        # ------------------- Single-value or other cases -------------------
        elif type == "monovalore" or type == "altro" or type == "media_annuale":
            continue
        
        # ------------------- Add chart to Word document -------------------
        if chart_created and chart_path:
            inline_shape = doc.add_picture(chart_path, width=Inches(6))
            inline = inline_shape._inline
            docPr = inline.docPr
            docPr.set('descr', 'Bar or boxplot chart visually representing the data values')
