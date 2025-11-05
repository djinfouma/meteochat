# ===========================
# Standard Chat Backend Script
# ===========================
# Handles dataset loading, Azure OpenAI setup, text chunking,
# vector storage creation, and retrieval system initialization
# for the standard (non-expert) chatbot.

import json
import os
import re
import io
import docx.oxml
from docx.oxml.ns import qn
import pandas as pd
from graphics_standard import generate_graphics, extract_values_month, extract_values_annual
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain_chroma import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import DirectoryLoader
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from flask import send_file

# ===========================
# Global state and chat tracking
# ===========================

chat_history = []# Stores user questions and model answers

# Need during the creation of the Document
conversation_text_precipitation = []
conversation_text_temperature = []
conversation_text_pressure = []

# ===========================
# Data loading and preprocessing
# ===========================

# Data loading from directories
loader_prec = DirectoryLoader('output/precipitazioni/')
loader_temp = DirectoryLoader('output/temperature/')
loader_pres = DirectoryLoader('output/pressione/')

# Load all files into memory
data_prec = loader_prec.load()
data_temp = loader_temp.load()
data_pres = loader_pres.load()

# Combine all data
all_data = data_prec + data_temp + data_pres


# ===========================
# Text chunking for vectorization
# ===========================

# Split documents into overlapping chunks for embeddings
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=50,
    length_function=len,
    is_separator_regex=False
)
splitted_data = text_splitter.split_documents(all_data)

# ===========================
# Load API credentials
# ===========================

# Load Azure OpenAI credentials from the secrets file
with open('secrets.json') as f:
    secrets = json.load(f)
secrets = secrets['gpt-4o']

# Initialize the embedding model
embeddings = AzureOpenAIEmbeddings(
    azure_endpoint=secrets['endpoint'],
    api_key=secrets['api_key']
)

# ===========================
# Chroma database setup
# ===========================

# Define Chroma vector store configuration
collection_metadata = {
    "hnsw:space": "cosine",
    "hnsw:M": 32,
    "hnsw:ef_construction": 128
}

# Check if the persistent database exists, else create it
if not os.path.exists('db'):
    store = Chroma.from_documents(
        splitted_data,
        embeddings,
        ids=[f"{item.metadata.get('source', 'unknown-source')}-{index}" for index, item in enumerate(splitted_data)],
        persist_directory='db',
        collection_metadata=collection_metadata
    )
    store.persist()
else:
    # Load existing Chroma database
    store = Chroma(persist_directory='db', embedding_function=embeddings)

# ===========================
# Retriever configuration
# ===========================

# Compute document count and optimize retrieval parameters
num_docs = len(store.get()["ids"])
k_value = max(10, num_docs // 4)  # Retrieve at least 10 docs, up to 25% of all

# Create a retriever with MMR (Maximal Marginal Relevance)
retriever = store.as_retriever(
    search_type="mmr", 
    search_kwargs={"k": k_value, "fetch_k": num_docs}
)

# ===========================
# Prompt
# ===========================
# Here, the system will define an improved PromptTemplate
# for the standard chatbot to ensure simplified, readable responses.
prompt = PromptTemplate(
    template="""
    You are a meteorologist who explains environmental data to a general audience.
    Your goal is to transform technical information into short, engaging, and clear narratives that highlight meaningful trends or changes.

    **Principles to follow:**
    - Tone: conversational, informative, and vivid — but not exaggerated.
    - Focus on clarity and insight more than storytelling flair.
    - Use simple metaphors or imagery only if they help understanding (avoid overly poetic language).
    - Keep answers concise and fact-driven.
    - Connect the data to real-world implications or everyday experience when possible.
    - Avoid technical jargon and excessive numbers — summarize trends in plain language.

    **Available context:**
    {context}
    - Year: the year
    - Month: the month
    - Mean Precipitation (mm): average precipitation for the month
    - Max Precipitation (mm): maximum precipitation in a single day during the month
    - Min Precipitation (mm): minimum precipitation in a single day during the month
    - Mode Precipitation (mm): most frequent precipitation value for the month
    - Mean Temperature (°C): average monthly temperature
    - Max Temperature (°C): maximum monthly temperature
    - Min Temperature (°C): minimum monthly temperature
    - Mode Temperature (°C): most frequent monthly temperature value
    - Mean Pressure (mbar): average atmospheric pressure for the month
    - Max Pressure (mbar): highest atmospheric pressure recorded during the month
    - Min Pressure (mbar): lowest atmospheric pressure recorded during the month
    - Mode Pressure (mbar): most frequent atmospheric pressure value for the month

    If the question is about precipitation, focus on precipitation-related values.
    If the question is about temperature, focus on temperature-related values.
    If the question is about pressure, focus on pressure-related values.

    When performing calculations, explain them briefly and simply.

    Answer the following question in a natural, informative way:
    {question}
    """,
    input_variables=['context', 'question'],
)

# ===========================
# AI Model Configuration
# ===========================
# Initializes the Azure OpenAI chat model and builds a RetrievalQA chain
# to handle standard weather-related queries.

# AI model setup
llm = AzureChatOpenAI(
    azure_endpoint=secrets['endpoint'],
    api_key=secrets['api_key'],
    api_version="2023-12-01-preview",
    deployment_name=secrets['deployment_name'],
    model="model_fine-tuned"
)

# Combine retriever and model into a QA pipeline
llm_chain = RetrievalQA.from_chain_type(
    llm=llm,
    chain_type='stuff',
    retriever=retriever,
    chain_type_kwargs={'prompt': prompt},
    return_source_documents=False
)

# -------------------------------------
# SUPPORT FUNCTION: EXTRACTS OR NORMALIZES NUMERIC DATA
# -------------------------------------
def normalize_time_series_df(df: pd.DataFrame, question_text: str = "") -> pd.DataFrame:
    """
    Normalizes a monthly or annual DataFrame for chart generation.
    Compatible with temperature, precipitation, and pressure datasets.
    Keeps all 12 months even if some values are missing (NaN).
    """
    if df is None or not isinstance(df, pd.DataFrame):
        return pd.DataFrame()
    if df.empty:
        return df.copy()

    df = df.copy()
    df.columns = [c.strip() for c in df.columns]

    df = df.loc[:, ~df.columns.duplicated()]

     # --- Identify month column ---
    month_col = next((c for c in df.columns if "month" in c.lower()), None)
    if not month_col:
        month_col = df.columns[0]  # assume first column if no 'month' found

    # --- Identify numeric columns ---
    numeric_cols = []
    for col in df.columns:
        if col == month_col:
            continue
        # Consider numeric if at least one value can be parsed as a number
        if pd.to_numeric(df[col], errors="coerce").notna().any():
            numeric_cols.append(col)

    if not numeric_cols:
        numeric_cols = [c for c in df.columns if c != month_col]

    # --- Sort by month ---
    month_order_full = ["January","February","March","April","May","June",
                        "July","August","September","October","November","December"]
    
    def month_to_index(m):
        m_str = str(m).strip().lower()
        # Handle numeric months (1–12)
        if m_str.isdigit() and 1 <= int(m_str) <= 12:
            return int(m_str)
        # Handle abbreviated month names
        abbrev_map = {m[:3].lower(): i+1 for i, m in enumerate(month_order_full)}
        return abbrev_map.get(m_str[:3], 99)   # unknown months go to the end

    df["month_index"] = df[month_col].apply(month_to_index)
    df = df.sort_values("month_index").drop(columns=["month_index"]).reset_index(drop=True)

    # --- Ensure all 12 months are present ---
    full_month_df = pd.DataFrame({month_col: month_order_full})
    df = full_month_df.merge(df, on=month_col, how="left")

    # Make sure column names are unique after merge
    seen = {}
    new_columns = []
    for c in df.columns:
        if c not in seen:
            seen[c] = 0
            new_columns.append(c)
        else:
            seen[c] += 1
            new_columns.append(f"{c}_{seen[c]}")
    df.columns = new_columns

    # Clean up duplicates and spaces
    df.columns = [c.strip() for c in df.columns]
    df = df.loc[:, ~df.columns.duplicated()]

    return df




def parse_markdown_table(md_text: str) -> pd.DataFrame:
    """
    Converts a Markdown table into a clean pandas DataFrame.
    Handles missing values and empty cells gracefully.
    Keeps all month rows even if some data are missing.
    """
    if not md_text or not isinstance(md_text, str):
        return pd.DataFrame()
    
    # Extract lines containing Markdown table delimiters
    lines = [l.strip() for l in md_text.splitlines() if "|" in l and not l.strip().startswith("|-") and not l.strip().startswith("```")]
    if len(lines) < 2:
        return pd.DataFrame()
    
    # Header row
    headers = [h.strip() for h in lines[0].strip("|").split("|") if h.strip()]
    
    data = []
    for line in lines[1:]:
        if "---" in line:
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        # Adjust cell count to match header count
        if len(cells) < len(headers):
            cells += ["" for _ in range(len(headers) - len(cells))]
        elif len(cells) > len(headers):
            cells = cells[:len(headers)]
        data.append(cells)
    
    df = pd.DataFrame(data, columns=headers)
    
    # Replace common missing data symbols with NaN
    df.replace({"-": pd.NA, "—": pd.NA, "–": pd.NA, "NA": pd.NA, "": pd.NA}, inplace=True)
    numeric_cols = [c for c in df.columns if c.lower() != "month"]
    # Convert numeric columns safely
    for col in numeric_cols:
        def parse_value(val):
            try:
                return float(val)
            except (ValueError, TypeError):
                return pd.NA
        df[col] = df[col].apply(parse_value)

    
    return df


def generate_numeric_for_graph(response_text, chat_history, llm):
    """
    Extracts a Markdown table from the model's response and converts it
    into a numeric DataFrame suitable for generating visual charts.
    Uses the model to reconstruct complete monthly data if partial info is found.
    """
    all_data_text = "\n".join([doc.page_content for doc in all_data])
    extraction_prompt = f"""
You are a data analyst specialized in environmental time series.
Extract a complete table of numeric values relevant to {chat_history[-1]["question"]}.
If the question refers to a specific year (e.g., 2024), include ALL 12 months of that year.
If the question refers to multiple years, include all available yearly data.
If only partial information is mentioned in the text, complete the table logically from the context — do NOT invent numbers.

Return ONLY a Markdown table, with columns:
| Month | Mean Temperature (°C) | Max Temperature (°C) | Min Temperature (°C) |
(if temperature question),
or
| Month | Mean Precipitation (mm) | Max Precipitation (mm) | Min Precipitation (mm) |
(if precipitation question),
or
| Month | Mean Pressure (mbar) | Max Pressure (mbar) | Min Pressure (mbar) |
(if pressure question).

If data are not available, leave cells blank, but keep all months listed (January–December).
Do NOT add explanations, only the table.

Context:
{response_text}
Full dataset reference from ARPA Lazio:{all_data_text}
"""

    try:
        # Ask the model to extract numeric data only as a Markdown table
        numeric_summary = llm.invoke(
            "Return ONLY a Markdown table. Do NOT add explanations.\n\n" + extraction_prompt
        )
        md_text = numeric_summary.content if hasattr(numeric_summary, "content") else str(numeric_summary)
    except Exception as e:
        print(f"[ERROR] LLM extraction failed: {e}")
        return
     # Check for Markdown table presence
    if "|" not in md_text:
        print("[WARN] Nessuna tabella Markdown trovata, salto estrazione.")
        return
    print("Md Text:")
    print(md_text)
    df = parse_markdown_table(md_text)
    df = normalize_time_series_df(df, question_text=chat_history[-1].get("question", ""))

    print("Parsed DataFrame:")
    print(df)

    if df.empty:
        return

     # Add parsed DataFrame to the last chat entry for graph generation
    chat_history[-1]["numeric_df"] = df

# ------------------------
# MAIN CHAT FUNCTION
# ------------------------
def standard_chat(user_input: str) -> str:
    """Handles chat interaction with the standard model."""
    global chat_history
    if not user_input:
        return "No message provided"
    try:
        # Send user input to the model
        response = llm_chain.invoke(user_input)
        response_text = response if isinstance(response, str) else response.get("result", "No response available")
        response_text = response_text.replace("**", "").replace("###", "")
        chat_history.append({"question": user_input, "answer": response_text})
        # Attempt to extract numeric data for charts if not already present
        generate_numeric_for_graph(response_text, chat_history,llm)


        return response_text
    except Exception as e:
        return f"Error generating response: {e}"


def download_standard_report():
    """Generates and downloads the conversation report in DOCX format."""
    if not chat_history:
        return "No conversation available"
     # Combine the full conversation into a readable string
    conversation_text = "\n".join([f"Q: {entry['question']}\nA: {entry['answer']}" for entry in chat_history])
    
    # Generate title
    title_prompt = f"Generate an engaging title for this conversation(don't add the keyword title):\n{conversation_text}"
    try:
        title_response = llm_chain.invoke(title_prompt)
        title = title_response if isinstance(title_response, str) else title_response.get("result", "Climate Report")
    except Exception as e:
        print(f"Error generating title: {e}")
        title = "Climate Report"

    # Generate summary
    summary_prompt = f"Imagine you are a meteorologist writing a story-like summary of this conversation for a general audience (3-4 sentences):\n{conversation_text}"
    try:
        summary_response = llm_chain.invoke(summary_prompt)
        summary = summary_response if isinstance(summary_response, str) else summary_response.get("result", "Summary unavailable.")
    except Exception as e:
        print(f"Error generating summary: {e}")
        summary = "Summary unavailable."

    # Generate keywords
    keywords_prompt = f"Generate 3-4 keywords that best summarize this conversation:\n{conversation_text}"
    try:
        keywords_response = llm_chain.invoke(keywords_prompt)
        keywords = keywords_response if isinstance(keywords_response, str) else keywords_response.get("result", "Keywords unavailable.")
    except Exception as e:
        print(f"Error generating keywords: {e}")
        keywords = "Keywords unavailable."

    # Helper function: get heading style (multilingual
    def get_heading_style(doc, level=1):
        """Returns the correct Word heading style ('Heading' or 'Titolo') based on language."""
        possible_names = [f"Heading {level}", f"Titolo {level}"]
        for name in possible_names:
            try:
                return doc.styles[name]
            except KeyError:
                continue
        raise ValueError(f"Nessuno stile trovato per livello {level}")


    # Create the Word document
    doc = Document()

     # Set global font
    doc.styles['Normal'].font.name = 'Century Gothic'

    # === HEADER LOGO CONFIGURATION ===
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]

     # Insert logo into the header 
    run = paragraph.add_run()
    run.add_picture('static/LogoMeteoChat_nero.png', width=Inches(1))  
    paragraph.alignment = 1 

    # Accessible green color palette for headings
    green_palette = {
        1: RGBColor(0x00, 0x6A, 0x4E),  # medium dark green
        2: RGBColor(0x00, 0x64, 0x00),  # dark green
        3: RGBColor(0x22, 0x8B, 0x22),  # forest green  
    }
    
    font_sizes = {
        1: Pt(20),  # main title
        2: Pt(16),  # secondary heading
        3: Pt(14),  # subheading
    }

    # Apply consistent heading font, color, and size
    for lvl in range(1, 4):
        try:
            style = get_heading_style(doc, level=lvl)
            style.font.name = 'Century Gothic'
            style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Century Gothic') 
            style.font.size = font_sizes[lvl]
            style.font.color.rgb = green_palette[lvl]
        except ValueError:
            pass
    # Helper: ensures all headings use the same font
    def set_heading_font(paragraph, font_name="Century Gothic"):
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # -------------------------------------
    # Add title, abstract, and keywords
    # -------------------------------------
    doc.core_properties.title = title
    title_set= doc.add_heading(title, level=1)
    set_heading_font(title_set)
    abstract_set = doc.add_heading("Abstract", 2)
    set_heading_font(abstract_set)
    doc.add_paragraph(summary)
    doc.add_paragraph(keywords)
    doc.add_page_break()

    # -------------------------------------
    # Add all Q&A pairs from the conversation
    # -------------------------------------
    for entry in chat_history:
        question_set = doc.add_heading(entry["question"], level=3)
        set_heading_font(question_set)
        doc.add_paragraph(entry["answer"])
        # Generate charts from extracted numeric data
        generate_graphics(doc, [entry])

    doc.add_page_break()
   
    
    # Generate conclusion
    conclusion_prompt = f"Eriting a conclusion for a general audience about this conversation (3-4 sentences):\n{conversation_text}"
    try:
        conclusion_response = llm_chain.invoke(conclusion_prompt)
        conclusion = conclusion_response if isinstance(conclusion_response, str) else conclusion_response.get("result", "Conclusion unavailable.")
    except Exception as e:
        print(f"Error generating conclusion: {e}")
        conclusion = "Conclusion unavailable."

    conclusion_set=doc.add_heading("Conclusion", 2)
    set_heading_font(conclusion_set)
    doc.add_paragraph(conclusion)

    # Save file
    file_path = "report_standard.docx"
    doc.save(file_path)
    return send_file(file_path, as_attachment=True)


