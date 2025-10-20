# ===========================
# Expert Chat Backend Script
# ===========================
# Handles data loading, AI model configuration, conversation management, 
# and report generation with charts and Word export.

import json
import os
import re
import io
import markdown2
import docx.oxml
from docx.oxml.ns import qn
import pandas as pd
from graphics import generate_graphics
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain_chroma import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import DirectoryLoader
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from flask import send_file


# ===========================
# Global variables and chat state
# ===========================

chat_history = []  # Stores conversation entries (Q&A pairs)

# Lists for grouped questions by data type
conversation_text_precipitation = []
conversation_text_temperature = []
conversation_text_pressure = []

# ===========================
# Load and preprocess datasets
# ===========================

# Load documents from subdirectories
loader_prec = DirectoryLoader('output/precipitazioni/')
loader_temp = DirectoryLoader('output/temperature/')
loader_pres = DirectoryLoader('output/pressione/')

# Load the text data
data_prec = loader_prec.load()
data_temp = loader_temp.load()
data_pres = loader_pres.load()

# Combine all loaded data
all_data = data_prec + data_temp + data_pres

# Split documents into smaller chunks for embedding
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=50,
    length_function=len,
    is_separator_regex=False
)
splitted_data = text_splitter.split_documents(all_data)

# ===========================
# Azure OpenAI Configuration
# ===========================

# Load API credentials from secrets.json
with open('secrets.json') as f:
    secrets = json.load(f)
secrets = secrets['gpt-4o']

# Initialize embeddings
embeddings = AzureOpenAIEmbeddings(
    azure_endpoint=secrets['endpoint'],
    api_key=secrets['api_key']
)

# Indexing configuration for ChromaDB
collection_metadata = {
    "hnsw:space": "cosine",
    "hnsw:M": 32,
    "hnsw:ef_construction": 128
}

# Check for database existence and initialize or load
if not os.path.exists('db'):
    store = Chroma.from_documents(
        splitted_data,
        embeddings,
        ids=[f"{item.metadata.get('source', 'unknown-source')}-{index}" for index, item in enumerate(splitted_data)],
        persist_directory='db',
        collection_metadata=collection_metadata
    )
    store.persist()
else:
    store = Chroma(persist_directory='db', embedding_function=embeddings)

# Adjust retrieval size dynamically based on dataset
num_docs = len(store.get()["ids"])
k_value = max(10, num_docs // 4)  # Retrieve 25% of total documents, min 10

# Setup retriever with MMR (diversity-based) search
retriever = store.as_retriever(
    search_type="mmr",
    search_kwargs={"k": k_value, "fetch_k": num_docs}
)

# ===========================
# Prompt Template for the AI model
# ===========================

prompt = PromptTemplate(
    template="""Consider {context} where:
    - Year: the year
    - Month: the month
    - Mean Precipitation (mm): average precipitation for the month
    - Max Precipitation (mm): maximum precipitation in a single day during the month
    - Min Precipitation (mm): minimum precipitation in a single day during the month
    - Mode Precipitation (mm): most frequent precipitation value for the month
    - Mean Temperature (°C): average monthly temperature
    - Max Temperature (°C): maximum monthly temperature
    - Min Temperature (°C): minimum monthly temperature
    - Mode Temperature (°C): most frequent monthly temperature value for the month
    - Mean Pressure (mbar): average atmospheric pressure for the month
    - Max Pressure (mbar): highest atmospheric pressure recorded during the month
    - Min Pressure (mbar): lowest atmospheric pressure recorded during the month
    - Mode Pressure (mbar): most frequent atmospheric pressure value for the month

    If the question is about precipitation, use the relevant data.  
    If the question is about temperature, use the temperature-related values.  
    If the question is about atmospheric pressure, use the pressure-related values.

    Always format any list of data as a **table** with clear headers.  
    If any data value is missing, unknown, or NaN, do NOT include that cell or row in the table.
    Never write "NaN", "None", or "null" — simply leave the value out or skip the row
    Example transformation:
    - January: 3.5 mm
    - February: 7.2 mm

    Should be formatted as:

    | Month    | Max Precipitation (mm) |
    |----------|------------------------|
    | January  | 3.5                    |
    | February | 7.2                    |

    When performing calculations:
    - **Do NOT use LaTeX, \[...\], $...$, \frac, or any math markup.**
    - Show calculations clearly as plain text, step by step.
    - Example:
        Average Pressure 2019 = (1011.43 + 1021.14 + 1018.56 + 1013.03 + ...) / 12  
        = 12175.59 / 12 = 1014.63 mbar
    - Use section labels such as:
        Data:
        Calculations:
        Conclusion:

    The final answer must be clear, structured, and easy to read in plain text format,  
    so that it can be rendered correctly in HTML or Word without LaTeX rendering.

    Answer the following question: {question}""",
    input_variables=['context', 'question'],
)

# ===========================
# Model setup
# ===========================

llm = AzureChatOpenAI(
    azure_endpoint=secrets['endpoint'],
    api_key=secrets['api_key'],
    api_version="2023-12-01-preview",
    deployment_name=secrets['deployment_name'],
    model="model_fine-tuned"
)

# Combine retriever and model into a single QA chain
llm_chain = RetrievalQA.from_chain_type(
    llm=llm,
    chain_type='stuff',
    retriever=retriever,
    chain_type_kwargs={'prompt': prompt},
    return_source_documents=False
)

# ===========================
# Functions
# ===========================

all_dataframes = []  # Stores extracted tables across the chat


def expert_chat(user_input: str) -> str:
    """Handles the expert chat interaction, converts Markdown tables to HTML, and saves conversation state."""
    global chat_history, latest_dataframe, all_dataframes
    
    if not user_input:
        return "No message provided"
    
    try:
        # Invoke the AI model
        response = llm_chain.invoke(user_input)
        response_text = response if isinstance(response, str) else response.get("result", "No response available")
        
        # Clean Markdown syntax for readability
        clean_text = response_text.replace("**", "").replace("###", "")
        all_tables = []

        # Extract all Markdown tables
        matches = re.findall(r"(\|.+\|[\s\S]+?)(?=\n\s*\n|\Z)", clean_text)
        for table_md in matches:
            # Convert Markdown table to DataFrame
            df = pd.read_csv(pd.io.common.StringIO(table_md), sep="|", engine="python").dropna(axis=1, how="all")
            df = df.apply(lambda x: x.str.strip() if x.dtype == "object" else x)
            df = df[~df.iloc[:, 0].str.fullmatch(r'-+') & df.iloc[:, 0].notna()]
            for col in df.columns[1:]:
                df[col] = pd.to_numeric(df[col], errors="coerce")

            latest_dataframe = df
            all_dataframes.append(df)
            all_tables.append(df)

            # Convert DataFrame to HTML for chat rendering
            html_table = df.to_html(index=False, border=0)
            html_table = f"""
<br></br>
<div style="overflow-x:auto;">
<style>
table {{ border-collapse: collapse; width: 100%; font-family: Arial, sans-serif; }}
th, td {{ border: 1px solid #aaa; padding: 8px; text-align: left; }}
th {{ background-color: #f2f2f2; }}
td {{ background-color: white; }}
</style>
{html_table}
</div>
<br></br>
"""
            clean_text = clean_text.replace(table_md, html_table)

        # Split text and tables, convert text blocks into paragraphs
        parts = re.split(r'(<div style="overflow-x:auto;">[\s\S]+?</div>)', clean_text)
        processed_parts = []
        for part in parts:
            if part.startswith('<div style="overflow-x:auto;">'):
                processed_parts.append(part)
            else:
                paragraphs = [p.strip() for p in part.split('\n\n') if p.strip()]
                for p in paragraphs:
                    processed_parts.append(f'<p>{p}</p>')

        clean_text_html = '\n'.join(processed_parts)

        # Save Q&A entry to chat history
        chat_history.append({
            "question": user_input,
            "answer_text": response_text,
            "answer_html": clean_text_html,
            "dataframes": all_tables
        })

        return clean_text_html

    except Exception as e:
        return f"Error generating response: {e}"


# Clear lists in case of multiple report generations
conversation_text_precipitation.clear()
conversation_text_temperature.clear()
conversation_text_pressure.clear()


def download_expert_report():
    """Generates and downloads a structured Word report from the full expert chat."""
    if not chat_history:
        return "No conversation available"

    # Combine conversation into plain text
    conversation_text = "\n".join([f"Q: {entry['question']}\nA: {entry['answer_text']}" for entry in chat_history])

    # --- Generate report metadata (title, summary, keywords) ---
    title_prompt = f"Generate a title for this conversation (don't add the word 'title'):\n{conversation_text}"
    try:
        title_response = llm_chain.invoke(title_prompt)
        title_text = title_response if isinstance(title_response, str) else title_response.get("result", "Climate Report")
        title = title_text.replace("**", "").replace("###", "")
    except Exception as e:
        print(f"Error generating title: {e}")
        title = "Climate Report"

    # Abstract / summary
    summary_prompt = f"Imagine you are a meteorologist tasked with writing an abstract for a report based on this conversation (3-4 sentences):\n{conversation_text}"
    try:
        summary_response = llm_chain.invoke(summary_prompt)
        summary_text = summary_response if isinstance(summary_response, str) else summary_response.get("result", "Summary unavailable.")
        summary = summary_text.replace("**", "").replace("###", "")
    except Exception as e:
        print(f"Error generating summary: {e}")
        summary = "Summary unavailable."

    # Keywords
    keywords_prompt = f"Generate 3–4 keywords that best summarize this conversation:\n{conversation_text}"
    try:
        keywords_response = llm_chain.invoke(keywords_prompt)
        keywords_text = keywords_response if isinstance(keywords_response, str) else keywords_response.get("result", "Keywords unavailable.")
        keywords = keywords_text.replace("**", "").replace("###", "")
    except Exception as e:
        print(f"Error generating keywords: {e}")
        keywords = "Keywords unavailable."

    # Helper: robust heading style finder (multilingual)
    def get_heading_style(doc, level=1):
        possible_names = [f"Heading {level}", f"Titolo {level}"]
        for name in possible_names:
            try:
                return doc.styles[name]
            except KeyError:
                continue
        raise ValueError(f"No heading style found for level {level}")

    # --- Create Word Document ---
    doc = Document()
    doc.styles['Normal'].font.name = 'Century Gothic'

    # Insert logo in header
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('static/LogoMeteoChat_nero.png', width=Inches(1))
    paragraph.alignment = 1

    # Define accessible green palette and heading styles
    green_palette = {
        1: RGBColor(0x00, 0x6A, 0x4E),
        2: RGBColor(0x00, 0x64, 0x00),
        3: RGBColor(0x22, 0x8B, 0x22),
    }
    font_sizes = {1: Pt(20), 2: Pt(16), 3: Pt(14)}

    for lvl in range(1, 4):
        try:
            style = get_heading_style(doc, level=lvl)
            style.font.name = 'Century Gothic'
            style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Century Gothic')
            style.font.size = font_sizes[lvl]
            style.font.color.rgb = green_palette[lvl]
        except ValueError:
            pass

    def set_heading_font(paragraph, font_name="Century Gothic"):
        """Applies font to all runs in a heading for consistent style."""
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # Add title, abstract, and keywords
    doc.core_properties.title = title
    title_set = doc.add_heading(title, level=1)
    set_heading_font(title_set)
    abstract_set = doc.add_heading("Abstract", 2)
    set_heading_font(abstract_set)
    doc.add_paragraph(summary)
    doc.add_paragraph(keywords)
    doc.add_page_break()

    # ===========================
    # Group chat by topic
    # ===========================
    for entry in chat_history:
        q_lower = entry["question"].lower()
        if re.search(r"\bprecipitation(s)?\b", q_lower):
            conversation_text_precipitation.append(entry)
        elif re.search(r"\btemperature(s)?\b", q_lower):
            conversation_text_temperature.append(entry)
        elif re.search(r"\bpressure(s)?\b", q_lower):
            conversation_text_pressure.append(entry)

    # --- Helper functions for inserting tables and graphs into the document ---
    def split_text_and_markdown_tables(text):
        """Splits text into blocks separating Markdown tables."""
        pattern = r'(\|.*\|\n\|[\-\s\|]*\n(?:\|.*\|\n?)*)'
        parts = re.split(pattern, text)
        return parts

    def contains_markdown_table(block):
        """Checks whether a text block contains a Markdown table."""
        return bool(re.match(r'\|.*\|', block.strip()))

    def add_question_answer_with_graph(doc, entry):
        """Inserts question, answer, tables, and charts for each entry."""
        question_text = entry["question"]
        question = question_text.replace("**", "").replace("###", "")
        answer_text = entry["answer_text"]
        answer = answer_text.replace("**", "").replace("###", "")
        dataframes_for_question = entry["dataframes"]

        question_set = doc.add_heading(question, level=3)
        set_heading_font(question_set)

        blocks = split_text_and_markdown_tables(answer)
        df_index = 0

        # Insert Word tables or paragraphs
        for block in blocks:
            if contains_markdown_table(block) and df_index < len(dataframes_for_question):
                df_corrente = dataframes_for_question[df_index]
                df_index += 1

                table = doc.add_table(rows=1, cols=len(df_corrente.columns))
                table.style = "Table Grid"

                for idx, col_name in enumerate(df_corrente.columns):
                    table.rows[0].cells[idx].text = str(col_name)

                for _, row in df_corrente.iterrows():
                    row_cells = table.add_row().cells
                    for idx, value in enumerate(row):
                        row_cells[idx].text = str(value)
            else:
                clean_block = block.strip()
                if clean_block:
                    doc.add_paragraph(clean_block)

        # Generate related graphs for the entry
        if dataframes_for_question:
            generate_graphics(doc, [{"question": question, "answer": answer}], dfs=dataframes_for_question)
        else:
            generate_graphics(doc, [{"question": question, "answer": answer}], dfs=None)

    # ===========================
    # Add all conversation sections
    # ===========================
    if conversation_text_precipitation:
        precipitation_set = doc.add_heading("Precipitation", level=2)
        set_heading_font(precipitation_set)
        for entry in conversation_text_precipitation:
            add_question_answer_with_graph(doc, entry)

    if conversation_text_temperature:
        temperature_set = doc.add_heading("Temperature", level=2)
        set_heading_font(temperature_set)
        for entry in conversation_text_temperature:
            add_question_answer_with_graph(doc, entry)

    if conversation_text_pressure:
        pressure_set = doc.add_heading("Pressure", level=2)
        set_heading_font(pressure_set)
        for entry in conversation_text_pressure:
            add_question_answer_with_graph(doc, entry)

    # ===========================
    # Summary Table and Conclusion
    # ===========================

    doc.add_page_break()
    data_set = doc.add_heading("Data", level=2)
    set_heading_font(data_set)

    # Prompt to generate summary table
    table_prompt = (
        "Generate a summary table in CSV format based on the following conversation. "
    "Each row must have exactly two columns: "
    "1) A concise keyword or phrase summarizing the question (e.g., 'highest temperature'), "
    "2) The numeric result from the answer, including contextual information (e.g., 'July, 37.553°C'). "
    "If multiple questions are present, create one row per question. "
    "Do not include any headers or extra text.\n"
    f"{conversation_text}"
    )
    try:
        table_response = llm_chain.invoke(table_prompt)
        table_text = table_response if isinstance(table_response, str) else table_response.get("result", "Table not available")
    except Exception as e:
        print(f"Error generating the table: {e}")
        table_text = "Table not available"

    if table_text and table_text != "Table not available":
        doc.add_paragraph("Summary Table:")
        
        # Split the AI response into rows and process them correctly
        rows = [row.strip() for row in table_text.split("\n") if row.strip()]
        
        # Create the table with headers
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Question"
        hdr_cells[1].text = "Numerical Data"
        
        # Set row cells as headers for accessibility
        for cell in table.rows[0].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(docx.oxml.parse_xml(r'<w:tcHeadr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
      
        
        for row in rows:
            cols = [col.strip() for col in row.split(",")]
            if len(cols) >= 2:
                row_cells = table.add_row().cells
                row_cells[0].text = cols[0]
                row_cells[1].text = ",".join(cols[1:])
    else:
        doc.add_paragraph("Table not available.")

    # Generate conclusion
    conclusion_prompt = f"Generate the conclusion for the report based on this conversation (3-4 sentences):\n{conversation_text}"
    try:
        conclusion_response = llm_chain.invoke(conclusion_prompt)
        conclusion = conclusion_response if isinstance(conclusion_response, str) else conclusion_response.get("result", "Conclusion unavailable.")
    except Exception as e:
        print(f"Error generating conclusion: {e}")
        conclusion = "Conclusion unavailable."

    conclusion_set = doc.add_heading("Conclusion", 2)
    set_heading_font(conclusion_set)
    doc.add_paragraph(conclusion)

    # Save and return file
    file_path = "report_expert.docx"
    doc.save(file_path)
    return send_file(file_path, as_attachment=True)
